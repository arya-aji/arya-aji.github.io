import docx
from docx.shared import Pt, Inches, RGBColor

# Create Document 1 for B2B Flow
doc1 = docx.Document()

title1 = doc1.add_heading('Alur Pemesanan dan Pembayaran Layanan Web Development (B2B)', 0)
title1.alignment = 1

p_info1 = doc1.add_paragraph()
p_info1.add_run('Merchant: ').bold = True
p_info1.add_run('Arya Aji Kusuma (aryaaji.com)')

doc1.add_heading('Langkah 1: Klien Memilih Layanan di Website', level=2)
p1 = doc1.add_paragraph()
p1.add_run('Keterangan: ').bold = True
p1.add_run('Klien mengunjungi halaman "Layanan & Produk" (https://aryaaji.com/pricing), melihat detail fitur, harga (dalam IDR), serta Syarat & Ketentuan. Klien kemudian menekan tombol "Minta Penawaran" atau "Chat di WhatsApp".\n\n')
run1 = p1.add_run('[HAPUS TEKS INI DAN MASUKKAN SCREENSHOT HALAMAN PRICING aryaaji.com DI SINI]')
run1.bold = True
run1.font.color.rgb = RGBColor(255, 0, 0)

doc1.add_heading('Langkah 2: Diskusi Kebutuhan dan Harga via Chat', level=2)
p2 = doc1.add_paragraph()
p2.add_run('Keterangan: ').bold = True
p2.add_run('Merchant dan klien mendiskusikan spesifikasi teknis dan menyepakati harga final dari layanan tersebut via WhatsApp atau Email.\n\n')
run2 = p2.add_run('[HAPUS TEKS INI DAN MASUKKAN SCREENSHOT CHAT WA/EMAIL DI SINI]')
run2.bold = True
run2.font.color.rgb = RGBColor(255, 0, 0)

doc1.add_heading('Langkah 3: Pembuatan Invoice / Payment Link via Midtrans', level=2)
p3 = doc1.add_paragraph()
p3.add_run('Keterangan: ').bold = True
p3.add_run('Setelah sepakat, Merchant login ke Dashboard Midtrans untuk membuat Payment Link sesuai dengan jumlah nominal yang disepakati.\n\n')
run3 = p3.add_run('[HAPUS TEKS INI DAN MASUKKAN SCREENSHOT DASHBOARD MIDTRANS SAAT MEMBUAT PAYMENT LINK DI SINI]')
run3.bold = True
run3.font.color.rgb = RGBColor(255, 0, 0)

doc1.add_heading('Langkah 4: Klien Menerima Link dan Membayar', level=2)
p4 = doc1.add_paragraph()
p4.add_run('Keterangan: ').bold = True
p4.add_run('Klien mengklik link tersebut, lalu memilih metode pembayaran yang disediakan oleh Midtrans. Transaksi selesai.\n\n')
run4 = p4.add_run('[HAPUS TEKS INI DAN MASUKKAN SCREENSHOT HALAMAN PEMBAYARAN SNAP MIDTRANS DI SINI]')
run4.bold = True
run4.font.color.rgb = RGBColor(255, 0, 0)

doc1.save('c:/Shared/Coding/arya-aji.github.io/Alur_Transaksi_B2B_Payment_Link.docx')

# Create Document 2 for SaaS Flow
doc2 = docx.Document()

title2 = doc2.add_heading('Alur Pembelian Produk SaaS - Siap Kuliah UI (B2C)', 0)
title2.alignment = 1

p_info2 = doc2.add_paragraph()
p_info2.add_run('Merchant: ').bold = True
p_info2.add_run('Arya Aji Kusuma (tryout.aryaaji.com)')

doc2.add_heading('Langkah 1: Pilihan Produk', level=2)
s1 = doc2.add_paragraph()
s1.add_run('Keterangan: ').bold = True
s1.add_run('User mengunjungi web aplikasi (https://tryout.aryaaji.com) dan memilih Paket Kredit yang tersedia pada tabel harga.\n\n')
run_s1 = s1.add_run('[HAPUS TEKS INI DAN MASUKKAN SCREENSHOT TABEL HARGA SKUI DI SINI]')
run_s1.bold = True
run_s1.font.color.rgb = RGBColor(255, 0, 0)

doc2.add_heading('Langkah 2: Checkout', level=2)
s2 = doc2.add_paragraph()
s2.add_run('Keterangan: ').bold = True
s2.add_run('User mengklik tombol beli dan sistem menampilkan form checkout/login di dalam web SKUI.\n\n')
run_s2 = s2.add_run('[HAPUS TEKS INI DAN MASUKKAN SCREENSHOT FORM CHECKOUT DI WEB SKUI DI SINI]')
run_s2.bold = True
run_s2.font.color.rgb = RGBColor(255, 0, 0)

doc2.add_heading('Langkah 3: Pembayaran via Snap API', level=2)
s3 = doc2.add_paragraph()
s3.add_run('Keterangan: ').bold = True
s3.add_run('Muncul popup pembayaran Midtrans Snap di layar secara langsung, dan user memilih metode pembayaran yang diinginkan.\n\n')
run_s3 = s3.add_run('[HAPUS TEKS INI DAN MASUKKAN SCREENSHOT POPUP SNAP MIDTRANS DI SINI]')
run_s3.bold = True
run_s3.font.color.rgb = RGBColor(255, 0, 0)

doc2.save('c:/Shared/Coding/arya-aji.github.io/Alur_Transaksi_SaaS_Snap.docx')
